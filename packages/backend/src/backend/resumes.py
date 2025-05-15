import copy
from docx.text.paragraph import Paragraph


from backend.docx_functions import (
    copy_paragraph_format,
    delete_paragraph,
    set_list_indent_level,
)
from backend.tailoring.schema import SerializedParagraph, SerializedRun


def add_runs_to_paragraph(
    paragraph: Paragraph,
    runs_data: list[SerializedRun],
    run_template=None,
):
    for run in runs_data:
        if run_template:
            new_run = copy.deepcopy(run_template)
            new_run.text = run.text
            paragraph._p.append(new_run._element)

        else:
            new_run = paragraph.add_run(run.text)

        if run.styles is not None:
            new_run.bold = "bold" in run.styles
            new_run.italic = "italic" in run.styles
            new_run.underline = "underline" in run.styles

    return paragraph


def find_next_non_anchor(paragraphs, lookup, start_idx=0):
    """
    Finds the first paragraph that does not match the lookup
    """

    idx = start_idx
    for para in paragraphs[start_idx:]:
        if para.text not in lookup:
            return para, idx
        idx += 1

    return paragraphs[-1], idx


def update_resume_section(
    section_content: list[Paragraph], new_paragraphs: list[SerializedParagraph]
):
    """
    Updates the section with new paragraphs

    Understanding this function:
        Anchor paragraphs are paragraphs that are the same between the original and the updated resume (often headings, dates or job position)
        We use them as "anchor" points to insert new paragraphs after
        So when we reach a new paragraph that is in the anchor_points_lookup, we find that same paragraph in the original resume
        and begin adding the following new paragraphs after it
        Note: we don't actually insert directly after the anchor paragraph, we instead assign a "pointer" paragraph as the
        next non-anchor paragraph
        The reason for this is so that we can use the insert_paragraph_before function, and insert our paragraphs in order
        This also makes it easier to use the styles from the pointer paragraph which will tend to be closer to the style we want for the new
        paragraphs as the anchor paragraphs often are headings and have the associated different styles
    """

    preserved_paragraphs = set([p.get_text() for p in new_paragraphs if p.preserved])
    anchor_paragraphs_lookup = {
        para.text: idx
        for idx, para in enumerate(section_content)
        if para.text in preserved_paragraphs
    }

    pointer_paragraph = section_content[0]
    pointer_paragraph_idx = 0
    for updated_para_raw in new_paragraphs:
        # If the paragraph is preserved, we don't need to do anything with it
        # Just move the pointer location to the next non-anchor
        if updated_para_raw.get_text() in anchor_paragraphs_lookup:
            # print("Found in lookup: ", updated_para.get_text()[:25])
            existing_paragraph_idx = anchor_paragraphs_lookup.get(
                updated_para_raw.get_text()
            )
            pointer_paragraph, pointer_paragraph_idx = find_next_non_anchor(
                section_content,
                anchor_paragraphs_lookup,
                start_idx=existing_paragraph_idx + 1,
            )
            continue

        # Otherwise, we need to insert a new paragraph
        new_para = pointer_paragraph.insert_paragraph_before(
            style=pointer_paragraph.style
        )
        copy_paragraph_format(pointer_paragraph, new_para)

        new_para = add_runs_to_paragraph(
            paragraph=new_para,
            runs_data=updated_para_raw.runs,
            run_template=pointer_paragraph.runs[0] if pointer_paragraph.runs else None,
        )

        # if updated_para_raw.list_indent_level is not None:
        set_list_indent_level(
            new_para,
            updated_para_raw.list_indent_level,
        )

        # Ideally, we move the pointer to the next paragraph if it is not an anchor
        # This allows us to maximally use the styles provided by the original document
        # However when we run out of lines with styling from the doc, we just keep using the last one's styling
        next_non_anchor, next_non_anchor_idx = find_next_non_anchor(
            section_content,
            anchor_paragraphs_lookup,
            start_idx=pointer_paragraph_idx + 1,
        )
        if next_non_anchor_idx == pointer_paragraph_idx + 1:
            pointer_paragraph, pointer_paragraph_idx = (
                next_non_anchor,
                next_non_anchor_idx,
            )
    # Delete all paragraphs that are not in the new paragraphs
    # These are just the old paragraphs from the original resume
    # So we delete them
    for para in section_content:
        if para.text not in anchor_paragraphs_lookup:
            delete_paragraph(para)


# if __name__ == "__main__":
#     # resumePath = fetch_resume("testUserId", "V3 Compressed Fabric.docx")
#     # resumePath = f"{RESUMES_PATH}/testUserId/V3Resume.docx"
#     resumePath = f"{RESUMES_PATH}/testUserId/Senior-Product-Manager-Resume-Example.docx"
#     parse_resume_for_sections(resumePath)
