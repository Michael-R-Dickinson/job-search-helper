import copy
import itertools
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.shared import OxmlElement
from docx.shared import Inches


def preprocess_paragraph(para):
    """
    Preprocess the text by removing extra whitespace.
    """
    clean_paragraph_whitespace(para)

    return para


def clean_paragraph_whitespace(para: Paragraph) -> None:
    """
    Removes unnecessary whitespace from the paragraph - IN PLACE

    Specifically:
    Collapse unnecessary whitespace between single-character runs (letters)
    while preserving styling on those runs, and normalize multiple
    whitespace runs into a single space.
    """
    runs = para.runs
    n = len(runs)

    # 1) Remove spaces between single-character runs
    for i in range(1, n - 1):
        prev_text = runs[i - 1].text
        curr_text = runs[i].text
        next_text = runs[i + 1].text

        # if this run is only whitespace, and neighbors are each exactly one non-space char
        if (
            curr_text.strip() == ""
            and len(prev_text.strip()) == 1
            and len(next_text.strip()) == 1
        ):
            runs[i].text = ""

    # 2) Normalize contiguous whitespace runs into a single space
    i = 0
    while i < len(runs):
        # detect runs that contain only whitespace
        if runs[i].text.strip() == "" and runs[i].text != "":
            # start of a whitespace-run group
            j = i
            # collect the group
            while j < len(runs) and runs[j].text.strip() == "" and runs[j].text != "":
                j += 1
            group = runs[i:j]
            # collapse group: first run → single space, others → empty
            if group:
                group[0].text = " "
                for extra in group[1:]:
                    extra.text = ""
            i = j
        else:
            i += 1


def is_builtin_heading_style(para):
    """True if the paragraph uses a built-in Heading style."""
    name = para.style.name or ""
    return name.startswith("Heading")


def compute_average_font_size(doc):
    """
    Scan all runs with a defined font size and return the average in Pt.
    """
    sizes = [
        run.font.size.pt
        for para in doc.paragraphs
        for run in para.runs
        if run.font.size
    ]
    return sum(sizes) / len(sizes) if sizes else None


def is_above_average_font_size(para, doc, multiplier=1.0):
    """
    True if any run in para has font.size > avg_size * multiplier.
    """

    # Compute avg font size once per doc
    avg_size = getattr(doc, "_avg_font_size", None)
    if avg_size is None:
        avg_size = compute_average_font_size(doc)
        doc._avg_font_size = avg_size

    if avg_size is None:
        return False
    threshold = avg_size * multiplier
    return any(run.font.size and run.font.size.pt > threshold for run in para.runs)


# TODO: Perhaps change this to any emphasized text (bold, underline)
def is_primarily_bold(para, bold_ratio=0.6):
    """
    True if more than bold_ratio of runs are bold.
    """
    runs = para.runs
    if not runs:
        return False

    bold_runs = sum(1 for r in runs if r.bold)
    if bold_runs / len(runs) >= bold_ratio:
        return True

    return False


def has_extra_whitespace(para, space_pt_threshold=6):
    """
    True if space_before or space_after ≥ space_pt_threshold (in points).
    """
    fmt = para.paragraph_format
    for attr in ("space_before", "space_after"):
        val = getattr(fmt, attr)
        pt = getattr(val, "pt", None)
        if pt is not None and pt >= space_pt_threshold:
            return True
    return False


def has_border(para):
    """
    True if the paragraph has any border defined in its pPr.
    """
    pPr = para._p.pPr
    if pPr is None:
        return False
    # look for any <w:pBdr> (paragraph borders)
    return pPr.find(qn("w:pBdr")) is not None


def is_mostly_uppercase(para, caps_ratio=0.8):
    """
    True if more than caps_ratio of total letters are uppercase.
    caps_ratio is the fraction of total letters that need to be uppercase for the paragraph to be considered uppercase.
    """

    text = para.text.strip()
    letters = [c for c in text if c.isalpha()]
    if letters:
        capital_letter_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
        return capital_letter_ratio >= caps_ratio

    return False


def is_title_case(para):
    """
    True if text is short (≤ max_words) AND
    each word starts uppercase (Title Case).
    """
    text = para.text.strip()
    words = text.split()
    return all(w[0].isupper() and w[1:].islower() for w in words if len(w) > 1)


def is_title_or_upper_case(para, caps_ratio=0.8):
    title_case = is_title_case(para)
    uppercase = is_mostly_uppercase(para, caps_ratio)
    return title_case or uppercase


SECTION_KEYWORDS = {
    "experience",
    "employment",
    "education",
    "skills",
    "projects",
    "certifications",
    "summary",
    "profile",
    "contact",
}


def has_keyword(para):
    """True if any SECTION_KEYWORDS appears in the paragraph text."""
    text = para.text.lower()
    return any(kw in text for kw in SECTION_KEYWORDS)


def has_few_distinct_words(para, max_words=4):
    """
    Return True if the paragraph has ≤ max_words distinct words.
    """
    text = para.text.lower()
    # extract word tokens (alphanumeric + underscore)
    words = re.findall(r"\b\w+\b", text)
    return len(set(words)) <= max_words


def is_likely_heading(para, doc):
    """
    Combine all heuristics. Return True if:
      - built-in heading style OR
      - keyword match
    OR at least two formatting signals of:
      font jump, bold/all-caps, whitespace, border, title/short.
    """
    if not para.text.strip():
        return False

    processed_para = copy.deepcopy(para)
    preprocess_paragraph(processed_para)

    signals = [
        is_builtin_heading_style(para),
        has_keyword(para),
        is_above_average_font_size(processed_para, doc),
        is_primarily_bold(processed_para),
        has_extra_whitespace(processed_para),
        has_border(processed_para),
        is_title_or_upper_case(processed_para),
    ]

    # The number of things making it likely a paragraph is a heading
    total_positive_indicators = sum(signals)

    # Headings should not be more than a few words long
    if not has_few_distinct_words(processed_para, max_words=4):
        total_positive_indicators -= 2

    is_heading = total_positive_indicators >= 2

    # ! Debugging
    # if is_heading:
    #     print(
    #         f"Indicators: {total_positive_indicators} \t Heading: {processed_para.text}"
    #     )

    return is_heading


def merge_identical_runs(para: Paragraph) -> None:
    """
    Merge consecutive runs in `para` that share exactly the same formatting,
    collapsing their text into a single run and removing the extras.
    """

    def run_format_key(run):
        """
        Create a key that uniquely identifies a run's full styling:
          - run.style.name
          - the raw <w:rPr> XML (empty string if none)
        """
        style_name = run.style.name or ""
        rPr = run._element.rPr  # the <w:rPr> element or None
        rPr_xml = rPr.xml if rPr is not None else ""
        return (style_name, rPr_xml)

    runs = para.runs
    if not runs:
        return

    # Group adjacent runs by identical formatting keys
    groups = []
    for key, group in itertools.groupby(runs, key=run_format_key):
        group = list(group)
        groups.append((key, group))

    # For each group with more than one run, merge them
    for _, group in groups:
        if len(group) <= 1:
            continue

        # Combine their text
        merged_text = "".join(r.text for r in group)

        # Assign merged text to the first run
        first_run = group[0]
        first_run.text = merged_text

        # Remove the extra run elements from the paragraph's XML
        for extra_run in group[1:]:
            para._p.remove(extra_run._element)


def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def iter_doc_paragraphs(doc: Document):
    for p in doc.iter_inner_content():
        if isinstance(p, Paragraph):
            yield p

        elif isinstance(p, Table):
            for row in p.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph


def get_list_level(para):
    """
    If `para` is a list item, returns its level as an int (0,1,2,…).
    Otherwise returns None.
    """
    pPr = para._p.pPr
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is None:
        return None
    return int(ilvl.get(qn("w:val")))


def set_list_indent_level(
    para, level: int, num_id: int = 1, indent_per_level=Inches(0.5)
):
    """
    Turn `para` into a list item at `level` (0=top,1=sub-item,2=sub-sub-item)
    referencing numbering definition `num_id`, and also set its left‐indent.

    - level: the w:ilvl value
    - num_id: w:numId (must exist in numbering.xml; usually 1 is the default list)
    - indent_per_level: how much to indent (inches) per nesting level
    """

    # Indent level 1 is actually the first sub-item (0 is the top level indent)
    level = level + 1

    p = para._p
    pPr = p.pPr or OxmlElement("w:pPr")
    if p.pPr is None:
        p.insert(0, pPr)

    old = pPr.find(qn("w:numPr"))
    if old is not None:
        pPr.remove(old)

    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

    para.paragraph_format.left_indent = indent_per_level * level


def copy_paragraph_format(src: Paragraph, tgt: Paragraph):
    src_pPr = src._p.pPr
    if src_pPr is None:
        return

    # Remove existing <w:pPr> from the target
    tgt_pPr = tgt._p.pPr
    if tgt_pPr is not None:
        tgt._p.remove(tgt_pPr)

    # Deep‐copy and insert the source <w:pPr>
    new_pPr = copy.deepcopy(src_pPr)
    tgt._p.insert(0, new_pPr)


if __name__ == "__main__":
    print("This module is not meant to be run directly.")
