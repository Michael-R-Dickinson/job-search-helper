from docx.text.paragraph import Paragraph


from docx_functions.modifications import (
    add_runs_to_paragraph,
    copy_paragraph_format,
    delete_paragraph,
    set_list_indent_level,
)
from LLM_tailoring.schema import SerializedParagraph


def is_anchor_paragraph(para: Paragraph, anchor_lookup: set[str]) -> bool:
    """
    Checks if the paragraph is an anchor paragraph based on whether it is in the lookup dictionary
    """
    return para.text in anchor_lookup


def find_next_non_anchor(paragraphs, anchor_lookup, start_idx=0):
    """
    Finds the first paragraph that does not match the lookup
    """

    idx = start_idx
    for para in paragraphs[start_idx:]:
        if not is_anchor_paragraph(para, anchor_lookup):
            return para, idx
        idx += 1

    return paragraphs[-1], idx


def add_paragraph_before(
    src_paragraph: Paragraph, new_paragraph_data: SerializedParagraph
) -> Paragraph:
    """
    Adds a new paragraph before the source paragraph
    """
    new_para = src_paragraph.insert_paragraph_before(style=src_paragraph.style)

    copy_paragraph_format(src_paragraph, new_para)

    new_para = add_runs_to_paragraph(
        paragraph=new_para,
        runs_data=new_paragraph_data.runs,
        run_template=src_paragraph.runs[0] if src_paragraph.runs else None,
    )

    set_list_indent_level(
        new_para,
        new_paragraph_data.list_indent_level,
    )

    return new_para


def update_resume_section(
    section_content: list[Paragraph], new_paragraphs: list[SerializedParagraph]
):
    """
    Updates the section with new paragraphs

    Understanding this function:
        Anchor paragraphs are paragraphs that are the same between the original and the updated resume (often headings, dates or job position)
        We use them as "anchor" points to insert new paragraphs after
        So when we reach a new paragraph that is in the anchor_points_lookup, we find that same paragraph in the original resume
        and begin adding the following new paragraphs after it
        Note: we don't actually insert directly after the anchor paragraph, we instead assign a "pointer" paragraph as the
        next non-anchor paragraph
        The reason for this is so that we can use the insert_paragraph_before function, and insert our paragraphs in order
        This also makes it easier to use the styles from the pointer paragraph which will tend to be closer to the style we want for the new
        paragraphs as the anchor paragraphs often are headings and have the associated different styles
    """

    preserved_paragraphs = set([p.get_text() for p in new_paragraphs if p.preserved])
    anchor_paragraphs_lookup = {
        para.text: idx
        for idx, para in enumerate(section_content)
        if para.text in preserved_paragraphs
    }

    pointer_paragraph = section_content[0]
    pointer_paragraph_idx = 0
    for updated_para_raw in new_paragraphs:
        # If the paragraph is preserved, we don't need to do anything with it
        # Just move the pointer location to the next non-anchor
        if updated_para_raw.get_text() in anchor_paragraphs_lookup:
            existing_paragraph_idx = anchor_paragraphs_lookup.get(
                updated_para_raw.get_text()
            )
            # Pointer paragraph is a paragraph in the ORIGINAL resume
            # And so the pointer_paragraph_idx should be used to index section_content
            pointer_paragraph, pointer_paragraph_idx = find_next_non_anchor(
                section_content,
                anchor_paragraphs_lookup,
                start_idx=existing_paragraph_idx + 1,
            )
            continue

        # Otherwise, we need to insert a new paragraph
        new_para = add_paragraph_before(
            src_paragraph=pointer_paragraph,
            new_paragraph_data=updated_para_raw,
        )

        # Ideally, we move the pointer to the next paragraph if it is not an anchor
        # This allows us to maximally use the styles provided by the original document
        # However because the last paragraph before an anchor often has different styling, we avoid
        # moving the pointer to the last paragraph before the anchor
        try:
            next_source_para_is_anchor = is_anchor_paragraph(
                section_content[pointer_paragraph_idx + 1], anchor_paragraphs_lookup
            )
            next_next_source_para_is_anchor = is_anchor_paragraph(
                section_content[pointer_paragraph_idx + 2], anchor_paragraphs_lookup
            )
            if not (next_source_para_is_anchor or next_next_source_para_is_anchor):
                pointer_paragraph, pointer_paragraph_idx = (
                    section_content[pointer_paragraph_idx + 1],
                    pointer_paragraph_idx + 1,
                )
        except IndexError:
            # We are at the end of the section, so we just keep the pointer where it is
            pass

    # Delete all paragraphs that are not in the new paragraphs
    # These are just the old paragraphs from the original resume
    # So we delete them
    for para in section_content:
        if para.text not in anchor_paragraphs_lookup:
            delete_paragraph(para)
