import copy
import itertools
import re
from docx import Document
from docx.oxml.ns import qn

from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.shared import OxmlElement
from docx.shared import Inches

from LLM_tailoring.resume.schema import SerializedRun
from docx_functions.paragraph_info import get_list_indent_level


def clean_paragraph_whitespace(para: Paragraph) -> None:
    """
    Removes unnecessary whitespace from the paragraph - IN PLACE

    Specifically:
    Collapse unnecessary whitespace between single-character runs (letters)
    while preserving styling on those runs, and normalize multiple
    whitespace runs into a single space.
    """
    runs = para.runs
    n = len(runs)

    # 1) Remove spaces between single-character runs
    for i in range(1, n - 1):
        prev_text = runs[i - 1].text
        curr_text = runs[i].text
        next_text = runs[i + 1].text

        # if this run is only whitespace, and neighbors are each exactly one non-space char
        if (
            curr_text.strip() == ""
            and len(prev_text.strip()) == 1
            and len(next_text.strip()) == 1
        ):
            runs[i].text = ""

    # 2) Normalize contiguous whitespace runs into a single space
    i = 0
    while i < len(runs):
        # detect runs that contain only whitespace
        if runs[i].text.strip() == "" and runs[i].text != "":
            # start of a whitespace-run group
            j = i
            # collect the group
            while j < len(runs) and runs[j].text.strip() == "" and runs[j].text != "":
                j += 1
            group = runs[i:j]
            # collapse group: first run → single space, others → empty
            if group:
                group[0].text = " "
                for extra in group[1:]:
                    extra.text = ""
            i = j
        else:
            i += 1


def remove_unnecessary_whitespace(text):
    """
    Collapse sequences of single‐character tokens into words, while preserving
    spacing between real words.

    E.g. "K E Y  S K I L L S" → "KEY SKILLS"
    """
    # 1) Split on two or more whitespace characters → these mark word boundaries
    segments = re.split(r"\s{2,}", text)
    cleaned_segments = []

    for seg in segments:
        # Normalize inner whitespace to single spaces, strip edges
        tokens = seg.strip().split()
        # If every token is exactly one character, collapse them into a word
        if tokens and all(len(tok) == 1 for tok in tokens):
            cleaned_segments.append("".join(tokens))
        else:
            # Otherwise treat as normal text, rejoin with single spaces
            cleaned_segments.append(" ".join(tokens))

    # Rejoin the segments with a single space (words were split by ≥2 spaces)
    return " ".join(cleaned_segments)


def clean_heading_text(text):
    """
    Cleans the heading text by removing unnecessary whitespace and
    converting text to lower case
    """

    cleaned_text = remove_unnecessary_whitespace(text)

    normalized_text = cleaned_text.lower()

    return normalized_text


def merge_identical_runs(para: Paragraph) -> None:
    """
    Merge consecutive runs in `para` that share exactly the same formatting,
    collapsing their text into a single run and removing the extras.
    """

    def run_format_key(run):
        """
        Create a key that uniquely identifies a run's full styling:
          - run.style.name
          - the raw <w:rPr> XML (empty string if none)
        """
        style_name = run.style.name or ""
        rPr = run._element.rPr  # the <w:rPr> element or None
        rPr_xml = rPr.xml if rPr is not None else ""
        return (style_name, rPr_xml)

    runs = para.runs
    if not runs:
        return

    # Group adjacent runs by identical formatting keys
    groups = []
    for key, group in itertools.groupby(runs, key=run_format_key):
        group = list(group)
        groups.append((key, group))

    # For each group with more than one run, merge them
    for _, group in groups:
        if len(group) <= 1:
            continue

        # Combine their text
        merged_text = "".join(r.text for r in group)

        # Assign merged text to the first run
        first_run = group[0]
        first_run.text = merged_text

        # Remove the extra run elements from the paragraph's XML
        for extra_run in group[1:]:
            para._p.remove(extra_run._element)


def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def clear_runs_only(para: Paragraph):
    """
    Remove only the <w:r> children of this paragraph, leaving its
    paragraph properties (styles, spacing, etc.) intact.
    """
    for run in list(para.runs):
        para._p.remove(run._element)


def delete_paragraphs(paragraphs: list[Paragraph]):
    for para in paragraphs:
        delete_paragraph(para)


def delete_paragraph(paragraph: Paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def add_runs_to_paragraph(
    paragraph: Paragraph,
    runs_data: list[SerializedRun],
    run_template=None,
):
    for run in runs_data:
        if run_template:
            new_run = copy.deepcopy(run_template)
            new_run.text = run.text
            paragraph._p.append(new_run._element)

        else:
            new_run = paragraph.add_run(run.text)

        if run.styles is not None:
            new_run.bold = "bold" in run.styles
            new_run.italic = "italic" in run.styles
            new_run.underline = "underline" in run.styles

    return paragraph


def set_list_indent_level(
    para, level: int | None, num_id: int = 1, indent_per_level=Inches(0.5)
):
    """
    Sets a paragraph to a bullet with indentation level - notably indent level 1 is the first level of a NESTED bullet
    so level 0 is the top level indent

    We have special rules for indent level = None
    - if level is None, we don't know whether the item should be indented once or not at all
      we just know it should not be a nested indent
    - so if the indent level is < 1 we make no changes
    - and if it is >= 1 we set the indent level to 0

    - level: the w:ilvl value
    - num_id: w:numId (must exist in numbering.xml; usually 1 is the default list)
    - indent_per_level: how much to indent (inches) per nesting level
    """

    # Indent level 1 is actually the first sub-item (0 is the top level indent)
    if level is None:
        current_level = get_list_indent_level(para)
        if current_level is None or current_level < 1:
            return
        else:
            level = 0

    level = level + 1
    p = para._p
    if p.pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    else:
        pPr = p.pPr

    old = pPr.find(qn("w:numPr"))
    if old is not None:
        pPr.remove(old)

    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

    para.paragraph_format.left_indent = indent_per_level * level


def copy_paragraph_format(src: Paragraph, tgt: Paragraph):
    src_pPr = src._p.pPr
    if src_pPr is None:
        return

    # Remove existing <w:pPr> from the target
    tgt_pPr = tgt._p.pPr
    if tgt_pPr is not None:
        tgt._p.remove(tgt_pPr)

    # Deep‐copy and insert the source <w:pPr>
    new_pPr = copy.deepcopy(src_pPr)
    tgt._p.insert(0, new_pPr)


if __name__ == "__main__":
    print("This module is not meant to be run directly.")
